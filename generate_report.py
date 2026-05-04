from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading_bold(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.bold = True

def create_detailed_report():
    doc = Document()

    # Title Section
    title = doc.add_heading('Project-Based Assessment Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph('Domain-Specific Generative AI Chatbot Using APIs')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True

    # Student Details
    doc.add_paragraph('Project Name: Campus Mate - The LPU Intelligence Hub')
    doc.add_paragraph('Domain: Education & University Administration (Campus Assistant)')
    doc.add_paragraph('Team Members:')
    doc.add_paragraph('1. Aman Chopra (Roll No: 59, Reg No: 12406294)')
    doc.add_paragraph('2. Aditya Thakur (Roll No: 26, Reg No: 12412223)')
    doc.add_paragraph('Guide/Faculty Name: _______________________')

    # 1. Introduction
    add_heading_bold(doc, '1. Introduction', 1)
    doc.add_paragraph(
        "In a large university ecosystem, students frequently face challenges navigating complex "
        "administrative systems, academic guidelines, and campus facilities. Finding immediate and "
        "accurate answers regarding UMS (University Management System), RMS (Request Management System), "
        "attendance, and placement rules is critical. Campus Mate is an innovative, domain-specific "
        "Generative AI chatbot designed to solve this real-world problem. Serving as an intelligent "
        "campus assistant, Campus Mate actively retrieves factual LPU guidelines to answer user queries. "
        "By leveraging the Groq Generative AI API (LLaMA 3.3 70B) integrated with a custom Retrieval-Augmented "
        "Generation (RAG) pipeline, Campus Mate provides a practical, responsible, and highly interactive "
        "AI solution that delivers a positive societal impact by drastically reducing administrative bottlenecks "
        "and fostering a smoother campus life for students."
    )

    # 2. Assessment Objective
    add_heading_bold(doc, '2. Assessment Objective', 1)
    doc.add_paragraph("This project successfully fulfills the core assessment objectives:")
    doc.add_paragraph("• Real-World Problem Identification: Addresses the lack of accessible, instant platforms for resolving student academic and administrative queries.")
    doc.add_paragraph("• Generative AI Application: Utilizes the Groq LLaMA 3.3 70B API to dynamically generate context-aware, highly accurate answers based on the user's input and matched dataset context.")
    doc.add_paragraph("• Technical Competence: Demonstrates a robust integration of RESTful APIs and machine learning embeddings within a custom-built, responsive Flask web application using Vanilla JavaScript.")
    doc.add_paragraph("• Design & Technical Decisions: Employs modern web technologies (Glassmorphism UI, real-time particle background animations) to create a seamless, engaging user experience.")

    # 3. Project Description
    add_heading_bold(doc, '3. Project Description', 1)
    doc.add_paragraph(
        "Campus Mate is tailored specifically for the Education and University Administration domain. "
        "Unlike general-purpose AI assistants (like standard ChatGPT) that might hallucinate university policies, "
        "Campus Mate is heavily constrained by system-level instructions and a strict RAG architecture to always "
        "rely on the verified dataset. It acts as an elite campus guide that retrieves specific protocols, "
        "identifies relevant departments (like DSW or Block 34), and provides step-by-step navigation for university "
        "portals, ensuring the interaction remains focused entirely on factual student assistance."
    )

    # 4. Scope and Requirements
    add_heading_bold(doc, '4. Scope and Requirements', 1)
    doc.add_paragraph("Mandatory Requirements Fulfilled:").bold = True
    doc.add_paragraph("• Clear Domain Definition: Education, specifically focused on campus administration and academic support.")
    doc.add_paragraph("• Generative AI API: Integrates Groq's LLaMA 3.3 70B API for high-speed natural language understanding and generation.")
    doc.add_paragraph("• Domain-Constrained Responses: System prompts strictly forbid the AI from using external knowledge, ensuring the output remains domain-relevant (LPU guidelines).")
    doc.add_paragraph("• Working Prototype: A fully functional, highly polished Flask web application featuring dynamic DOM manipulation and a responsive interface.")
    
    doc.add_paragraph("Optional Enhancements Implemented:").bold = True
    doc.add_paragraph("• Semantic Search Engine: Uses the all-MiniLM-L6-v2 transformer model to convert user questions into embeddings for accurate context matching.")
    doc.add_paragraph("• Keyword Boosting: A fail-safe mechanism that boosts specific highly-critical university keywords (RMS, UMS, Duty Leave) into the context.")
    doc.add_paragraph("• Advanced User Interface: Features a premium Glassmorphism design, ambient glow effects, quick-action chips, and responsive toast notifications.")

    # 5. Tools and Technologies
    add_heading_bold(doc, '5. Tools and Technologies', 1)
    doc.add_paragraph("• Generative AI API: Groq API (LLaMA 3.3 70B Versatile)")
    doc.add_paragraph("• Programming Languages: HTML5, CSS3, JavaScript (Vanilla ES6+), Python 3")
    doc.add_paragraph("• Frameworks/Libraries: Flask (Backend Server), Scikit-Learn (TF-IDF & Cosine Similarity), SentenceTransformers (Embeddings)")
    doc.add_paragraph("• Deployment/Hosting: Ready for Render / Heroku Web Application Hosting.")

    # 6. Data Collection and Domain Knowledge Preparation
    add_heading_bold(doc, '6. Data Collection and Domain Knowledge Preparation', 1)
    doc.add_paragraph(
        "To ensure Campus Mate functions effectively as a university assistant, domain knowledge regarding "
        "academic rules, placement criteria, and administrative processes was meticulously curated."
    )
    doc.add_paragraph("• Sources Studied: Over 1,590+ verified Q&A pairs were curated covering attendance rules, bonafide certificates, examination patterns (CA/MTE/ETE), and hostel regulations.")
    doc.add_paragraph("• Influence on Prompt Design: The gathered domain knowledge was used to engineer a robust System Prompt. Instead of a generic prompt, the system instruction explicitly commands the model to:\n"
                      "  1. Only use facts from the provided DATABASE CONTEXT.\n"
                      "  2. Provide a natural, conversational response using Markdown formatting.\n"
                      "  3. Use bullet points for sequences (e.g., UMS -> RMS -> DL).\n"
                      "  4. Refrain from hallucinating if information is missing from the context.")

    # 7. Model Configuration Awareness
    add_heading_bold(doc, '7. Model Configuration Awareness', 1)
    doc.add_paragraph(
        "Understanding model behavior parameters is critical for Campus Mate's success. When interfacing with the Groq API, the following configurations are heavily optimized:"
    )
    doc.add_paragraph("• Temperature (Creative vs. Factual): For a university assistant bot, a very low temperature (0.2) is ideal. Since the bot provides administrative policies and rules, creativity is not desired. A low temperature ensures arguments are deterministic, factual, and strictly aligned with the retrieved context without hallucinating non-existent university rules.")
    doc.add_paragraph("• Output Token Limit: Set to 800 tokens to ensure the responses remain concise, clear, and readable on mobile devices without cutting off abruptly.")
    doc.add_paragraph("• Semantic Fallback Threshold: The Cosine Similarity threshold is strictly monitored (0.20-0.25). If confidence is too low, the API call is bypassed entirely, and a fallback mechanism informs the student that the data is unavailable.")

    # Workflow Flowchart
    doc.add_page_break()
    add_heading_bold(doc, '8. System Architecture & Workflow Flowchart', 1)
    doc.add_paragraph("The following diagram illustrates the complete data flow from the user interface, through the semantic search pipeline, to the Generative AI API.")
    try:
        doc.add_picture('workflow_diagram.png', width=Inches(6.0))
    except Exception:
        doc.add_paragraph('[Workflow Diagram Screenshot Missing]')

    # Implementation Evidence
    doc.add_page_break()
    add_heading_bold(doc, '9. Implementation Evidence', 1)
    
    doc.add_heading('API Call Screenshot', level=2)
    doc.add_paragraph("The code snippet below demonstrates the successful integration of the Generative AI API (Groq) with the backend:")
    try:
        doc.add_picture('q12_api_call.png', width=Inches(6.0))
    except Exception:
        doc.add_paragraph('[API Call Screenshot missing]')

    doc.add_heading('Chatbot Working Interface', level=2)
    doc.add_paragraph("The user interface features a premium Glassmorphism design handling dynamic queries:")
    try:
        doc.add_picture('q13_chatbot_ui.png', width=Inches(6.0))
    except Exception:
        doc.add_paragraph('[Chatbot Working Interface Screenshot missing]')

    # GitHub Repository Link
    doc.add_heading('10. GitHub Repository Link', level=2)
    doc.add_paragraph('Repository URL: https://github.com/amanch0pra/campusmate')

    doc.save('Campus_Mate_Detailed_Report.docx')

if __name__ == '__main__':
    create_detailed_report()
